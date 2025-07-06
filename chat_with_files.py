import streamlit as st
import os
import pandas as pd
import fitz  # PyMuPDF
import extract_msg
import json
from bs4 import BeautifulSoup
import email
from email import policy
from email.parser import BytesParser
from docx import Document as DocxDocument
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.llms.openai import OpenAI
from tempfile import NamedTemporaryFile

def extract_text_from_file(uploaded_file):
    ext = uploaded_file.name.lower().split('.')[-1]
    content = ""

    if ext == "pdf":
        with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_file.read())
            doc = fitz.open(tmp.name)
            content = "\n".join(page.get_text() for page in doc)

    elif ext == "txt":
        content = uploaded_file.read().decode("utf-8")

    elif ext == "csv":
        df = pd.read_csv(uploaded_file)
        content = df.to_string(index=False)

    elif ext == "xlsx":
        df = pd.read_excel(uploaded_file, sheet_name=None)
        content = ""
        for name, sheet in df.items():
            content += f"Sheet: {name}\n"
            content += sheet.to_string(index=False) + "\n"

    elif ext == "msg":
        with NamedTemporaryFile(delete=False, suffix=".msg") as tmp:
            tmp.write(uploaded_file.read())
            msg = extract_msg.Message(tmp.name)
            content = f"Subject: {msg.subject}\nFrom: {msg.sender}\nTo: {msg.to}\nBody:\n{msg.body}"

    elif ext == "oft":
        content = "[OFT files not directly supported outside Outlook. Please convert to .msg or copy manually.]"

    elif ext == "docx":
        with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.read())
            doc = DocxDocument(tmp.name)
            content = "\n".join([para.text for para in doc.paragraphs])

    elif ext == "json":
        data = json.load(uploaded_file)
        content = json.dumps(data, indent=2)

    elif ext == "html":
        soup = BeautifulSoup(uploaded_file.read(), "html.parser")
        content = soup.get_text()

    elif ext == "eml":
        with NamedTemporaryFile(delete=False, suffix=".eml") as tmp:
            tmp.write(uploaded_file.read())
            with open(tmp.name, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
            content = f"Subject: {msg['subject']}\nFrom: {msg['from']}\nTo: {msg['to']}\nBody:\n{msg.get_body(preferencelist=('plain')).get_content()}"

    else:
        content = f"[Unsupported file type: {ext}]"

    return content

# Setup Streamlit
st.set_page_config(page_title="Chat with Any File", layout="wide")
st.title("📁 Chat with Any File")

# OpenAI Key from Streamlit secrets
os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"] if "OPENAI_API_KEY" in st.secrets else os.getenv("OPENAI_API_KEY")

uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "csv", "xlsx", "msg", "oft", "docx", "json", "html", "eml"], accept_multiple_files=True)

if uploaded_files:
    docs = []
    for file in uploaded_files:
        extracted_text = extract_text_from_file(file)
        docs.append(Document(text=extracted_text, metadata={"filename": file.name}))

    if st.button("🔍 Build Chatbot"):
        with st.spinner("Indexing documents..."):
            llm = OpenAI(model="gpt-3.5-turbo", temperature=0.1)
            Settings.llm = llm
            index = VectorStoreIndex.from_documents(docs)
            st.session_state.query_engine = index.as_query_engine()
            st.session_state.chat_history = []
        st.success("✅ Chatbot is ready!")

# Chat interface
if "query_engine" in st.session_state:
    query = st.chat_input("Ask a question about the uploaded files:")
    if query:
        with st.spinner("Thinking..."):
            response = st.session_state.query_engine.query(query)
            st.session_state.chat_history.append((query, response.response))

    for q, a in st.session_state.chat_history:
        with st.chat_message("user", avatar="👤"):
            st.markdown(q)
        with st.chat_message("assistant", avatar="🤖"):
            st.markdown(a)

    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🗑️ Clear Chat"):
            st.session_state.chat_history = []
    with col2:
        if st.download_button("📥 Download Chat", data="\n\n".join([f"User: {q}\nBot: {a}" for q, a in st.session_state.chat_history]),
                              file_name="chat_history.txt", mime="text/plain"):
            st.success("Chat history downloaded!")
